import io
import re
from typing import Optional
from fastapi import UploadFile
import PyPDF2
from docx import Document

class FileProcessor:
    """Service for extracting text from uploaded CV files"""
    
    async def extract_text(self, file: UploadFile) -> str:
        """
        Extract text from uploaded file (PDF, DOC, DOCX)
        """
        try:
            file_extension = file.filename.lower().split('.')[-1]
            content = await file.read()
            
            if file_extension == 'pdf':
                return self._extract_from_pdf(content)
            elif file_extension in ['doc', 'docx']:
                return self._extract_from_word(content)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            raise Exception(f"Error processing file: {str(e)}")
    
    def _extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF using PyPDF2 only"""
        try:
            text = ""
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return self._clean_text(text)
        except Exception as e:
            raise Exception(f"Failed to extract PDF text: {str(e)}")
    
    def _extract_from_word(self, content: bytes) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(io.BytesIO(content))
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return self._clean_text(text)
            
        except Exception as e:
            raise Exception(f"Failed to extract Word document text: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace and newlines
        text = re.sub(r'\n+', '\n', text)
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might interfere with processing
        text = re.sub(r'[^\w\s\-.,;:()\[\]/@#%&+]', '', text)
        
        return text.strip() 